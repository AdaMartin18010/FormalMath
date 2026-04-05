#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
多格式导出器
Multi-Format Exporter

支持Markdown、HTML、JSON、PDF、DOCX等多种格式导出
"""

import json
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime


class MultiFormatExporter:
    """
    多格式导出器
    
    将生成的文档导出为多种格式，便于不同场景使用
    """
    
    def __init__(self, output_dir: Path):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        self.export_dir = self.output_dir / "export"
        self.export_dir.mkdir(exist_ok=True)
    
    def export_all(self, source_dir: Path, format_type: str) -> Path:
        """
        导出所有文档为指定格式
        
        Args:
            source_dir: 源文档目录
            format_type: 导出格式 (markdown, html, json, pdf, docx)
            
        Returns:
            导出文件路径
        """
        source_dir = Path(source_dir)
        
        if format_type == "markdown":
            return self._export_markdown_bundle(source_dir)
        elif format_type == "html":
            return self._export_html_bundle(source_dir)
        elif format_type == "json":
            return self._export_json_bundle(source_dir)
        elif format_type == "pdf":
            return self._export_pdf_guide(source_dir)
        elif format_type == "docx":
            return self._export_docx_guide(source_dir)
        else:
            raise ValueError(f"不支持的格式: {format_type}")
    
    def _export_markdown_bundle(self, source_dir: Path) -> Path:
        """导出Markdown合集"""
        output_path = self.export_dir / "formalmath_complete.md"
        
        lines = [
            "# FormalMath 完整文档\n",
            f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n",
            "**包含内容**: API文档、概念图谱、Lean4定理\n\n",
            "---\n\n"
        ]
        
        # 收集所有markdown文件
        md_files = list(source_dir.rglob("*.md"))
        
        for md_file in sorted(md_files):
            try:
                content = md_file.read_text(encoding='utf-8')
                # 调整标题级别
                content = self._adjust_heading_levels(content, 1)
                lines.append(content)
                lines.append("\n\n---\n\n")
            except Exception as e:
                print(f"  警告: 读取 {md_file} 失败: {e}")
        
        output_path.write_text(''.join(lines), encoding='utf-8')
        print(f"  ✓ Markdown合集已导出: {output_path}")
        return output_path
    
    def _export_html_bundle(self, source_dir: Path) -> Path:
        """导出HTML合集"""
        output_path = self.export_dir / "formalmath_complete.html"
        
        css_styles = """
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { 
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; 
            line-height: 1.6; color: #333; background: white; 
            max-width: 900px; margin: 0 auto; padding: 40px 20px;
        }
        h1 { color: #1a73e8; border-bottom: 3px solid #1a73e8; padding-bottom: 10px; margin: 40px 0 20px; }
        h2 { color: #333; margin: 30px 0 15px; }
        h3 { color: #555; margin: 25px 0 10px; }
        p { margin: 10px 0; }
        code { background: #f5f5f5; padding: 2px 6px; border-radius: 3px; font-family: 'Consolas', monospace; }
        pre { background: #2d2d2d; color: #f8f8f2; padding: 20px; border-radius: 8px; overflow-x: auto; }
        table { width: 100%; border-collapse: collapse; margin: 20px 0; }
        th, td { border: 1px solid #ddd; padding: 12px; text-align: left; }
        th { background: #f8f9fa; font-weight: 600; }
        tr:nth-child(even) { background: #f8f9fa; }
        ul, ol { margin: 10px 0; padding-left: 30px; }
        blockquote { border-left: 4px solid #1a73e8; padding-left: 20px; margin: 20px 0; color: #555; }
        .page-break { page-break-before: always; }
        .toc { background: #f8f9fa; padding: 20px; border-radius: 8px; margin: 20px 0; }
        .toc a { color: #1a73e8; text-decoration: none; }
        @media print { body { padding: 0; } }
        """
        
        html_parts = [
            '<!DOCTYPE html>',
            '<html lang="zh-CN">',
            '<head>',
            '    <meta charset="UTF-8">',
            '    <title>FormalMath 完整文档</title>',
            f'    <style>{css_styles}</style>',
            '</head>',
            '<body>',
            '    <div class="header">',
            '        <h1>📚 FormalMath 完整文档</h1>',
            f'        <p>生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>',
            '    </div>',
            '    <div class="toc">',
            '        <h2>目录</h2>',
            '        <ul>',
            '            <li><a href="#api-docs">API文档</a></li>',
            '            <li><a href="#concept-graphs">概念图谱</a></li>',
            '            <li><a href="#lean4-docs">Lean4形式化</a></li>',
            '        </ul>',
            '    </div>'
        ]
        
        # 收集并转换Markdown内容
        md_files = list(source_dir.rglob("*.md"))
        
        for md_file in sorted(md_files)[:20]:  # 限制数量
            try:
                content = md_file.read_text(encoding='utf-8')
                html_content = self._markdown_to_html(content)
                html_parts.append(f'    <div class="page-break"></div>')
                html_parts.append(html_content)
            except Exception as e:
                print(f"  警告: 处理 {md_file} 失败: {e}")
        
        html_parts.extend([
            '</body>',
            '</html>'
        ])
        
        output_path.write_text('\n'.join(html_parts), encoding='utf-8')
        print(f"  ✓ HTML合集已导出: {output_path}")
        return output_path
    
    def _export_json_bundle(self, source_dir: Path) -> Path:
        """导出JSON数据包"""
        output_path = self.export_dir / "formalmath_data.json"
        
        data = {
            'metadata': {
                'generated_at': datetime.now().isoformat(),
                'version': '1.0.0',
                'source_dir': str(source_dir)
            },
            'documents': []
        }
        
        # 收集所有文档的元数据
        md_files = list(source_dir.rglob("*.md"))
        
        for md_file in md_files:
            try:
                content = md_file.read_text(encoding='utf-8')
                # 提取标题
                title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
                title = title_match.group(1) if title_match else md_file.stem
                
                doc_info = {
                    'path': str(md_file.relative_to(source_dir)),
                    'title': title,
                    'word_count': len(content),
                    'line_count': len(content.split('\n'))
                }
                data['documents'].append(doc_info)
            except Exception as e:
                pass
        
        output_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')
        print(f"  ✓ JSON数据包已导出: {output_path}")
        return output_path
    
    def _export_pdf_guide(self, source_dir: Path) -> Path:
        """生成PDF导出指南"""
        output_path = self.export_dir / "pdf_export_guide.md"
        
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        content = f'''# PDF 导出指南

**生成时间**: {timestamp}

## 推荐的PDF生成方法

### 方法1: 使用 Pandoc

```bash
# 安装 pandoc 和 wkhtmltopdf
# Windows: winget install pandoc
# macOS: brew install pandoc
# Ubuntu: sudo apt-get install pandoc wkhtmltopdf

# 导出为PDF
pandoc docs/generated/export/formalmath_complete.md \\
  -o docs/generated/export/formalmath.pdf \\
  --pdf-engine=xelatex \\
  -V CJKmainfont="SimSun" \\
  --toc \\
  --number-sections
```

### 方法2: 使用 Chrome/Edge 浏览器

1. 打开 `formalmath_complete.html`
2. 按 `Ctrl+P` (Windows) 或 `Cmd+P` (Mac)
3. 选择 "另存为PDF"
4. 设置页眉页脚和页码
5. 点击保存

### 方法3: 使用 Python 库

```python
# 安装依赖
# pip install markdown weasyprint

from markdown import markdown
from weasyprint import HTML, CSS

# 读取Markdown
with open('formalmath_complete.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# 转换为HTML
html_content = markdown(md_content, extensions=['tables', 'fenced_code'])

# 添加样式
css = """
@page {{ size: A4; margin: 2cm; }}
body {{ font-family: "Noto Sans CJK SC", sans-serif; }}
"""

# 生成PDF
HTML(string=html_content).write_pdf(
    'formalmath.pdf',
    stylesheets=[CSS(string=css)]
)
```

## PDF设置建议

| 设置项 | 推荐值 | 说明 |
|--------|--------|------|
| 页面大小 | A4 | 标准文档尺寸 |
| 页边距 | 2cm | 适合打印和阅读 |
| 字体 | Noto Sans CJK SC | 支持中文 |
| 代码块样式 | 等宽字体 | 保持代码格式 |
| 目录 | 启用 | 方便导航 |
| 页码 | 底部居中 | 标准格式 |

## 文档结构

生成的PDF将包含以下内容：

1. **封面** - 文档标题和元信息
2. **目录** - 自动生成的章节索引
3. **API文档** - 所有API参考
4. **概念图谱** - 数学概念及其关系
5. **Lean4形式化** - 定理和证明
6. **附录** - 补充材料

---
*由 FormalMath 多格式导出器创建*
'''
        
        output_path.write_text(content, encoding='utf-8')
        print(f"  ✓ PDF导出指南已生成: {output_path}")
        return output_path
    
    def _export_docx_guide(self, source_dir: Path) -> Path:
        """生成DOCX导出指南"""
        output_path = self.export_dir / "docx_export_guide.md"
        
        content = f'''# DOCX (Word) 导出指南

**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## 推荐的DOCX生成方法

### 方法1: 使用 Pandoc

```bash
# 导出为DOCX
pandoc docs/generated/export/formalmath_complete.md \\
  -o docs/generated/export/formalmath.docx \\
  --reference-doc=template.docx \\
  --toc
```

### 方法2: 使用 Python

```python
# 安装依赖
# pip install python-docx markdown

from docx import Document
from docx.shared import Inches, Pt
from markdown import markdown
from bs4 import BeautifulSoup

def markdown_to_docx(md_file, docx_file):
    # 读取Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换为HTML
    html_content = markdown(md_content)
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 创建Word文档
    doc = Document()
    
    # 处理HTML元素
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'pre']):
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'pre':
            # 代码块
            p = doc.add_paragraph()
            p.style = 'Code'
            p.add_run(element.text)
    
    # 保存
    doc.save(docx_file)

# 使用
markdown_to_docx(
    'formalmath_complete.md',
    'formalmath.docx'
)
```

### 方法3: 使用在线转换工具

1. 访问 [CloudConvert](https://cloudconvert.com/md-to-docx) 或类似网站
2. 上传 `formalmath_complete.md`
3. 选择输出格式为 DOCX
4. 下载转换后的文件

## Word样式设置建议

| 样式 | 字体 | 大小 | 其他 |
|------|------|------|------|
| 标题1 | 黑体 | 18pt | 段前段后12pt |
| 标题2 | 黑体 | 14pt | 段前段后6pt |
| 正文 | 宋体 | 12pt | 1.5倍行距 |
| 代码 | Consolas | 10pt | 灰色底纹 |

---
*由 FormalMath 多格式导出器创建*
'''
        
        output_path.write_text(content, encoding='utf-8')
        print(f"  ✓ DOCX导出指南已生成: {output_path}")
        return output_path
    
    def _adjust_heading_levels(self, content: str, base_level: int) -> str:
        """调整Markdown标题级别"""
        lines = content.split('\n')
        adjusted = []
        
        for line in lines:
            if line.startswith('#'):
                # 计算当前级别
                current_level = len(line) - len(line.lstrip('#'))
                new_level = min(current_level + base_level, 6)
                adjusted.append('#' * new_level + line[current_level:])
            else:
                adjusted.append(line)
        
        return '\n'.join(adjusted)
    
    def _markdown_to_html(self, content: str) -> str:
        """简单Markdown转HTML"""
        # 标题
        content = re.sub(r'^###\s+(.+)$', r'<h3>\1</h3>', content, flags=re.MULTILINE)
        content = re.sub(r'^##\s+(.+)$', r'<h2>\1</h2>', content, flags=re.MULTILINE)
        content = re.sub(r'^#\s+(.+)$', r'<h1>\1</h1>', content, flags=re.MULTILINE)
        
        # 代码块
        content = re.sub(
            r'```(\w+)?\n(.*?)\n```',
            r'<pre><code>\2</code></pre>',
            content,
            flags=re.DOTALL
        )
        
        # 行内代码
        content = re.sub(r'`([^`]+)`', r'<code>\1</code>', content)
        
        # 表格（简化处理）
        lines = content.split('\n')
        in_table = False
        html_lines = []
        
        for line in lines:
            if '|' in line and not in_table:
                in_table = True
                html_lines.append('<table>')
            
            if in_table:
                if '|' not in line:
                    in_table = False
                    html_lines.append('</table>')
                    html_lines.append(line)
                else:
                    cells = [c.strip() for c in line.split('|') if c.strip()]
                    if cells and not all(c.replace('-', '') == '' for c in cells):
                        html_lines.append('<tr>' + ''.join(f'<td>{c}</td>' for c in cells) + '</tr>')
            else:
                html_lines.append(line)
        
        if in_table:
            html_lines.append('</table>')
        
        content = '\n'.join(html_lines)
        
        # 段落
        paragraphs = content.split('\n\n')
        html_paragraphs = []
        
        for p in paragraphs:
            p = p.strip()
            if p and not p.startswith('<'):
                html_paragraphs.append(f'<p>{p}</p>')
            else:
                html_paragraphs.append(p)
        
        return '\n'.join(html_paragraphs)


def main():
    """主函数"""
    exporter = MultiFormatExporter(output_dir=Path("docs/generated"))
    
    source_dir = Path("docs/generated")
    
    for fmt in ["markdown", "html", "json", "pdf", "docx"]:
        try:
            output_path = exporter.export_all(source_dir, fmt)
            print(f"导出成功: {output_path}")
        except Exception as e:
            print(f"导出 {fmt} 失败: {e}")


if __name__ == '__main__':
    main()
